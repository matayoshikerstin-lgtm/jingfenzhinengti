import docx

doc = docx.Document(r'test_plan\经分智能体测试方案2.0.docx')

with open('debug_table1.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        if '示例' in p.text or '缺口' in p.text:
            f.write(f'Paragraph {i}: {p.text}\n')

    t = doc.tables[1]
    f.write("\nTable 1:\n")
    for r in t.rows:
        f.write(" | ".join([c.text.replace('\n', ' ') for c in r.cells]) + "\n")

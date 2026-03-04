import docx
import sys

def modify_manual_to_auto():
    file_path = r'test_plan\经分智能体测试方案2.0.docx'
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    replacements = {
        "手动对比版": "自动化对比+人工复检版",
        "直接手动提问/判断": "通过自动化/RPA提问并判断",
        "由测试人员在网页直接手动提问回复值，与Ground Truth（标准答案）对比判断": "通过自动化脚本获取回复，使用代码对比或大模型裁判与Ground Truth对比判断",
        "手动对比": "自动化对比",
        "测试人员将大模型的回复与Ground Truth的数值进行对比": "通过脚本自动将智能体的回复与Ground Truth的数值进行对比，对极差结果进行手工复检",
        "手动判断": "自动判断",
        "测试人员针对这些关键点（事实 + 逻辑 + 建议）分别打分给出得分理由": "大模型裁判针对这些关键点（事实 + 逻辑 + 建议）分别打分给出得分理由，对异常低分进行手工复检",
        "并手动判断准确性": "并通过自动化机制评估准确性",
        "由手动核对准确性": "由自动化脚本或大模型裁判核对准确性",
        "由测试人员比对Ground Truth": "基于自动化对比Ground Truth，对极差/异常结果进行人工复检"
    }

    for p in doc.paragraphs:
        original_text = p.text
        new_text = original_text
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
        
        # Apply changes if modified
        if new_text != original_text:
            # We must replace text while trying to keep formatting.
            # A simple way for whole paragraph replacement is to clear and add text, 
            # but if there are multiple runs, we can just clear and add a single run.
            style = p.style
            p.clear()
            p.add_run(new_text)

    try:
        doc.save(file_path)
        print("成功将'手动'替换为'自动化对比结合人工复检'等表述。")
    except Exception as e:
        print(f"保存文件失败: {e}")

if __name__ == '__main__':
    modify_manual_to_auto()
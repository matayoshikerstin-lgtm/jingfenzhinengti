import docx
from docx.shared import Pt, RGBColor
import sys

def insert_llm_judge():
    file_path = r'test_plan\经分智能体测试方案2.0.docx'
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # Find the paragraph "四. 测试指标" to insert before it
    insert_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith('四. 测试指标'):
            insert_idx = i
            break
            
    if insert_idx == -1:
        print("Could not find '四. 测试指标'")
        return
        
    # Insert new content before "四. 测试指标"
    # To insert before a paragraph in python-docx, we can use p.insert_paragraph_before()
    target_p = doc.paragraphs[insert_idx]
    
    p0 = target_p.insert_paragraph_before()
    r0 = p0.add_run('（新增）LLM-as-a-Judge 自动化打分法（推荐用于开放问题）')
    r0.bold = True
    r0.font.size = Pt(12)
    
    p1 = target_p.insert_paragraph_before('鉴于大模型（智能体）的输出为自由文本，且通常不会严格按照设定的“三段式”固定格式输出。为保证评分的客观性、一致性与高效性，测试方案引入 LLM-as-a-Judge（用大模型当裁判）的自动化评分机制。')
    
    p2 = target_p.insert_paragraph_before('核心评估理念：')
    p2.runs[0].bold = True
    target_p.insert_paragraph_before('评分时忽略智能体回答的排版格式（如是否带小标题），纯看其生成的语义是否涵盖了基准答案中的三大核心考核点：\n'
                                      '• 事实与数据层：是否找对底层数据，有无幻觉。\n'
                                      '• 归纳与分析层：是否把数据转化为业务洞察。\n'
                                      '• 落地与建议层：是否给出可执行的管理建议。')

    p3 = target_p.insert_paragraph_before('LLM 裁判指令 (Prompt) 模板：')
    p3.runs[0].bold = True
    
    prompt_text = """【系统设定】
你是一个资深的企业经营分析总监，现在你需要给实习生（经分智能体）生成的业务分析报告打分。

【输入信息】
1. 用户提问：[测试用例中的开放问题]
2. 参考基准答案：[测试方案中提供的带数据和建议的三段式标准答案]
3. 智能体实际回答：[被测经分智能体生成的实际长文本]

【打分规则（满分 5 分）】
请忽略智能体回答的排版格式，纯看语义是否涵盖了基准答案的核心考点：
- 准确性（2分）：实际回答中引用的业务数据、数值和客户名称，是否与基准答案一致，有无编造（幻觉）？
- 洞察力（1.5分）：实际回答是否指出了基准答案中提到的业务痛点或现象根因？
- 建议有效性（1.5分）：实际回答是否给出了与基准答案方向一致或同样具有业务指导意义的管理建议？

【输出格式】
1. 得分：[X / 5分]
2. 扣分理由：（如果满分则填“涵盖全面，数据精准”）"""

    p_prompt = target_p.insert_paragraph_before()
    run_prompt = p_prompt.add_run(prompt_text)
    run_prompt.font.name = 'Courier New'
    run_prompt.font.size = Pt(10)

    p4 = target_p.insert_paragraph_before('自动化测试实施路径：')
    p4.runs[0].bold = True
    target_p.insert_paragraph_before('1. 通过 RPA 脚本或 API 批量拉取智能体对 80 道开放题的实际回答，并落盘至 result.xlsx 的“智能体回复”列。\n'
                                     '2. 编写评测脚本，循环读取“问题 + 标准答案 + 智能体回复”，将其拼接到上述 Prompt 模板中。\n'
                                     '3. 调用更高阶的大模型（如 GPT-4o 或 Claude 3.5 Sonnet）接口，让其以裁判身份输出分数与扣分理由。\n'
                                     '4. 汇总分数，计算整体平均分以及“准确性、洞察力、建议”等细分维度的达标率。')
    
    target_p.insert_paragraph_before('') # Empty line

    # Now let's remove the Appendix we added previously
    # Find "附录：开放问题评估与打分标准（LLM-as-a-Judge 自动化打分法）" and delete everything from there to the end
    appendix_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if '附录：开放问题评估与打分标准' in p.text:
            appendix_idx = i
            break
            
    if appendix_idx != -1:
        # Delete paragraphs from appendix_idx to the end
        for p in doc.paragraphs[appendix_idx:]:
            p._element.getparent().remove(p._element)

    try:
        doc.save(file_path)
        print(f"成功将 LLM-as-a-Judge 修改到正文中，并删除了附录。")
    except Exception as e:
        print(f"保存文件失败: {e} (请确认文件未在 Word 中打开)")

if __name__ == '__main__':
    insert_llm_judge()
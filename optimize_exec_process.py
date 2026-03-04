import docx

def optimize_execution_process():
    file_path = r'test_plan\经分智能体测试方案2.0.docx'
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # Find the execution process section
    start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if '五. 测试执行流程' in p.text or '测试执行流程' in p.text:
            start_idx = i
            break
            
    if start_idx == -1:
        print("未找到测试执行流程章节")
        return

    # Clear old paragraphs in this section (we'll delete the next 4 paragraphs which are the old process)
    # Be careful not to delete beyond the document
    paragraphs_to_remove = []
    for i in range(start_idx, len(doc.paragraphs)):
        text = doc.paragraphs[i].text
        if i > start_idx and ('六.' in text or '七.' in text or text.strip() == ''):
            # stop if we hit the next section or empty lines after it
            if '六.' in text or '七.' in text:
                break
        paragraphs_to_remove.append(doc.paragraphs[i])

    # Insert new highly optimized paragraphs before the first removed paragraph
    target_p = doc.paragraphs[start_idx]
    
    p0 = target_p.insert_paragraph_before('五. 自动化与智能化测试执行流程')
    p0.style = target_p.style # try to keep heading style
    
    new_steps = [
        "1. 题库构建与基准锁定 (Test Case Engineering)：基于真实的「商机表」与「合同表」等底层核心数据，构建 80 道高拟真度经营分析问题（50 道高阶逻辑非开放题 + 30 道老板视角灵魂拷问开放题）。通过 Python 脚本计算并绝对锁定各题的 Ground Truth 及「三段式」评估基准，确保后续裁判标准的零误差与客观性。",
        
        "2. RPA 无人值守执行 (Robotic Process Automation)：引入 Playwright 自动化引擎，挂载 RPA 脚本对智能体交互界面进行无人值守跑批提问。脚本将自动注入测试数据、监听动态流式输出、提取完整回答文本，并将结果秒级落盘至数据表（result.xlsx），彻底消除人工操作带来的低效与人为干扰。",
        
        "3. 双轨智能交叉验证 (Dual-Track Validation)：\n"
        "   • 非开放题（硬性校验）：通过正则提取与实体识别脚本，将智能体回复的数值、列表与 Ground Truth 进行毫秒级精准比对。\n"
        "   • 开放题（LLM-as-a-Judge）：将智能体输出的非结构化长文本喂入高阶裁判大模型，严格按照「事实准确度、业务洞察力、建议可行性」三维评分矩阵进行语义解析与量化打分。仅对裁判给出的异常低分或涉嫌“严重幻觉”的案例触发人工复检（Human-in-the-loop），实现极高水准的评测自动化率。",
        
        "4. 缺陷洞察与高维报告产出 (Metrics & Insights Reporting)：聚合准确性、一致性与无幻觉性三大核心指标，自动生成可视化评测矩阵。拒绝只交出冰冷的测试分数，着重提炼智能体在“意图理解偏差、底层数据检索遗漏、深层逻辑推理断层”等维度的典型 Bad Case，用以反向倒逼经分智能体底层 RAG 链路与架构的精准迭代。"
    ]
    
    for step in new_steps:
        target_p.insert_paragraph_before(step)
        
    # Remove the old paragraphs
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    try:
        doc.save(file_path)
        print("【五. 测试执行流程】已进行极致优化！")
    except Exception as e:
        print(f"保存文件失败: {e}")

if __name__ == '__main__':
    optimize_execution_process()
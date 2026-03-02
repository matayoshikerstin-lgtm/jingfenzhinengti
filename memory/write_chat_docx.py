# 将聊天记忆写入 docx，运行: py memory/write_chat_docx.py
import os
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt

def main():
    doc = Document()
    doc.add_heading('聊天记录记忆整理', 0)
    doc.add_paragraph('整理自与 Cursor 的对话，便于换工作区后保留要点。项目：jingfenzhinengti（经分智能体测试）。')
    doc.add_paragraph()

    doc.add_heading('一、测试问题集建设', level=1)
    doc.add_paragraph(
        '依据：test_plan + test_data/total_data.xlsx（多 Sheet：商机表、合同表等）。'
        '非开放题 40 道：带基准答案 + 标注数据来源 Sheet（[商机表]/[合同表]）；题型含单表查询、聚合、排序/TOP N、多表关联。'
        '开放题：面向老板场景，字段均来自 total_data，带参考答案。'
        '构造规则：维度来自 test_plan，数据来自 total_data；题干用唯一键避免重名（如合同用合同编号）；开放题考虑“老板会问什么”。'
    )
    doc.add_paragraph()

    doc.add_heading('二、非开放题与基准答案修正（以当前 CSV/数据为准）', level=1)
    items = [
        ('Q4 华为合同金额', '合同名称有两条，改用合同编号 HT202410067 唯一确定；答案 41480000.0。'),
        ('Q26 需求明确阶段平均金额', '字段：销售流程=需求明确，商机金额 opp_amount；答案 2269.82（80 条）。Excel 筛选后求平均需用 SUBTOTAL(101, 列)，不能用 AVERAGE(范围)。'),
        ('Q27 第一季度预计签单总金额', '字段：预计签单日期 est_deal_date 在 2025-01-01～2025-03-31；汇总商机金额之和。'),
        ('Q31 商机金额前 3 名', '当前 fact_opportunity CSV 中第三名为招商银行股份有限公司业务中台项目 (4991.0)，非联想集团。'),
        ('Q34 净利润最低合同', '应为中国石油化工集团公司智能客服合同 (88190.66)，非“中国石化客户画像合同”。'),
        ('Q35 赢单率最高商机', '若有多条并列最高则取金额最大的一条。答案：华为技术有限公司AI大模型应用项目 (1.0)。'),
        ('Q37 华为合同按金额前 2 个', '客户/最终用户字段：final_user_name。前 2：① AI大模型应用合同 (41480000)；② 智能营销合同 (38950000)，第三才是智慧运营合同。'),
    ]
    for title, text in items:
        p = doc.add_paragraph()
        p.add_run(title + '：').bold = True
        p.add_run(' ' + text)
    doc.add_paragraph()

    doc.add_heading('三、字段与数据源对应', level=1)
    doc.add_paragraph('合同表：合同金额/销售总金额 sales_amount_total；净利润 net_profit；最终用户/客户 final_user_name；合同编号 contract_approval_code。')
    doc.add_paragraph('商机表：商机金额 opp_amount；预计签单日期 est_deal_date；销售流程 sales_process；预计赢单率 win_rate。')
    doc.add_paragraph('数据源：脚本基准答案来自 fact_opportunity_202602111036.csv、fact_contract_202602101728.csv；若用 total_data.xlsx 需以 xlsx 为准重算。')
    doc.add_paragraph()

    doc.add_heading('四、Excel 使用注意', level=1)
    doc.add_paragraph('筛选后只对可见行统计：用 SUBTOTAL。求和 SUBTOTAL(9 或 109, 列)；平均 SUBTOTAL(101, 列)；最大值 SUBTOTAL(104, 列)。排序时选“扩展选定区域”。')
    doc.add_paragraph()

    doc.add_heading('五、测试指标（评估智能体）', level=1)
    doc.add_paragraph('准确性：1～5 分专家均分，目标均≥4.5。一致性：20 次零差异率目标 100%。无幻觉性：无幻觉率目标 100%。')
    doc.add_paragraph()

    doc.add_heading('六、项目内其他', level=1)
    doc.add_paragraph('result.xlsx 乱码：用 fix_result_xlsx.py 转成真正 xlsx 输出 result_fixed.xlsx。周报：2 月 9～15 日个人小结。Cursor 聊天记录随工作区，换文件夹后可能看不到，重要内容建议保存到 memory 或文档。')
    doc.add_paragraph()
    doc.add_paragraph('生成时间：对话结束后整理。')

    out = os.path.join(os.path.dirname(__file__), 'chat_memory.docx')
    doc.save(out)
    print('已生成:', out)

if __name__ == '__main__':
    main()

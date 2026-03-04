import docx

def complete_optimize_plan():
    file_path = r'test_plan\经分智能体测试方案2.0.docx'
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    replacements = {
        "问题驱动测试，手动计算": "问题驱动测试，自动化评估与人工复检",
        "不使用裁判模型，直接手动计算/判断开放问题和非开放问题的值（准确性）。另外开放性问题待定。": "通过 RPA 自动化获取智能体回复，使用脚本精确对比非开放问题，并引入 LLM-as-a-Judge 对开放问题进行评分，异常结果触发人工复检。",
        "不使用裁判模型，": "采用自动化验证体系，",
        "非开放问题（客观，手动计算）": "非开放问题（客观事实类，脚本自动化比对）",
        "开放问题（主观，手动计算）": "开放问题（主观分析类，LLM-as-a-Judge 自动评分）",
        "手动检查事实": "大模型裁判核对事实",
        "手动检查逻辑": "大模型裁判验证逻辑",
        "手动检查建议": "大模型裁判评估建议",
        "手动检查完整": "大模型裁判判定完整度",
        "准确性手动计算（对不对）": "准确性自动评估",
        "手动检查输入数据/关系路径": "触发人工复检底层数据路径",
        "手动检查输入数据：漏数据支撑/具体行动（用户不满意）": "LLM 裁判判定为低分，触发人工复检缺陷明细",
        "自动判断准确性": "LLM裁判自动评估",
        "直接手动提问/判断": "通过自动化RPA提问并判断"
    }

    for p in doc.paragraphs:
        original_text = p.text
        new_text = original_text
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
        
        if new_text != original_text:
            style = p.style
            p.clear()
            p.add_run(new_text)

    # 检查表格中的文本并替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original_text = p.text
                    new_text = original_text
                    for old, new in replacements.items():
                        if old in new_text:
                            new_text = new_text.replace(old, new)
                    
                    if new_text != original_text:
                        p.clear()
                        p.add_run(new_text)

    try:
        doc.save(file_path)
        print("测试方案全面优化完成：已彻底清除'手动'相关过时表述，全量升级为自动化RPA+大模型裁判策略。")
    except Exception as e:
        print(f"保存文件失败: {e}")

if __name__ == '__main__':
    complete_optimize_plan()
